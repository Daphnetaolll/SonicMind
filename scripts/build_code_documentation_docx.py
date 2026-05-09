"""Build the polished SonicMind code documentation DOCX.

The Markdown version in docs/CODE_DOCUMENTATION.md is the lightweight source of
truth. This script creates the reader-facing Word document with consistent
heading, table, callout, and page furniture styles.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SonicMind_Code_Documentation.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(93, 109, 126)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=2)
    run = header.add_run("SonicMind Code Documentation | 2026-05-09")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0)
    run = footer.add_run("SonicMind | Engineering Reference")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=18)
    run = kicker.add_run("Engineering Reference Guide")
    set_run_font(run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=8)
    run = title.add_run("SonicMind Code Documentation")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=24)
    run = subtitle.add_run("Architecture, runtime modes, API boundaries, billing, retrieval, and verification map")
    set_run_font(run, size=13, color=DARK_BLUE)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [2700, CONTENT_WIDTH_DXA - 2700])
    rows = [
        ("Last updated", "2026-05-09"),
        ("Frontend", "https://sonicmind.onrender.com"),
        ("Backend API", "https://sonicmind-api.onrender.com"),
        ("Primary use", "Fast onboarding for engineers reviewing, deploying, or extending SonicMind"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=100, bottom=100)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        label_run = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(label_run, size=10, color=INK, bold=True)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    doc.add_page_break()


def add_paragraph(doc: Document, text: str, *, style: str | None = None, bold_label: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph)
    if bold_label and text.startswith(bold_label):
        run = paragraph.add_run(bold_label)
        set_run_font(run, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_label) :])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)


def add_callout(doc: Document, title: str, body: str, *, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.5, color=DARK_BLUE, bold=True)
    p.add_run("\n")
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)

    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=10, color=INK, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_data):
            set_cell_shading(cell, WHITE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.2)
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK)

    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = INK
    doc.add_paragraph()


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4)
        run = p.add_run(item)
        set_run_font(run, color=INK)


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4)
        run = p.add_run(item)
        set_run_font(run, color=INK)


def build_document() -> None:
    doc = Document()
    style_document(doc)
    add_running_header_footer(doc)
    add_cover(doc)

    doc.add_heading("1. Purpose and Reader Path", level=1)
    add_paragraph(
        doc,
        "This guide explains the current SonicMind codebase after the migration to a deployed React frontend and FastAPI backend. It is designed for engineers who need to review, deploy, debug, or extend the project without reverse-engineering every file first.",
    )
    add_callout(
        doc,
        "Current deployed surface",
        "Frontend: https://sonicmind.onrender.com | Backend API: https://sonicmind-api.onrender.com | Health: https://sonicmind-api.onrender.com/api/health",
    )

    doc.add_heading("2. High-Level Code Map", level=1)
    add_table(
        doc,
        ["Area", "Path", "Responsibility"],
        [
            ["React app", "frontend/src/", "Browser UI, routing, chat rendering, pricing actions, Spotify cards."],
            ["API client", "frontend/src/api/client.js", "Shared Axios client, bearer token injection, and API helpers."],
            ["Auth state", "frontend/src/store/authStore.js", "Persisted bearer token, user, chat turns, latest response, and UI settings."],
            ["FastAPI app", "backend/main.py", "Routes, auth dependency, CORS, health, account status, billing, and chat."],
            ["API schemas", "backend/schemas.py", "Pydantic request/response contracts for browser-safe APIs."],
            ["Billing service", "backend/services/billing_service.py", "Stripe Checkout, Portal, direct plan changes, webhook processing."],
            ["Chat service", "backend/services/chat_service.py", "Quota check, question logging, answer generation, usage charging."],
            ["Runtime settings", "src/settings.py", "Production lightweight vs semantic FAISS mode resolution."],
            ["RAG pipeline", "src/rag_pipeline.py", "Query routing, evidence gathering, answer synthesis, Spotify planning."],
            ["Database schema", "src/db/schema.py", "PostgreSQL tables, indexes, and repeatable schema setup."],
            ["Tests", "tests/ and frontend/tests/e2e/", "Pytest coverage and production browser smoke tests."],
        ],
        [1800, 2650, 4910],
    )

    doc.add_heading("3. Chat Request Flow", level=1)
    add_code_block(
        doc,
        [
            "React ChatPage",
            "-> frontend API client POST /api/chat",
            "-> FastAPI auth dependency",
            "-> chat_service.answer_user_question",
            "-> quota_service.get_quota_status",
            "-> question_repository creates question log",
            "-> rag_pipeline answers the question",
            "-> Spotify resolution runs when useful and configured",
            "-> successful answer records usage",
            "-> response returns answer, sources, Spotify cards, and plan usage",
        ],
    )
    add_bullet_list(
        doc,
        [
            "Quota is checked before answering and usage is deducted only after a successful text answer.",
            "Spotify failures should not break the main answer path.",
            "The response includes updated plan and quota fields so the frontend can refresh counters immediately.",
        ],
    )

    doc.add_heading("4. Account Status and Plan Source of Truth", level=1)
    add_paragraph(
        doc,
        "GET /api/me returns the current user plus quota and feature state. The displayed plan is derived from quota/source-of-truth state rather than trusting a stale users.plan row. This is important because Stripe webhook reconciliation can discover paid access after checkout or subscription updates.",
    )
    add_table(
        doc,
        ["Response area", "Meaning"],
        [
            ["user", "Browser-safe user fields including id, email, display name, plan, and subscription status."],
            ["usage", "Current plan, feature limits, remaining questions, and extra credit state."],
            ["plan display", "Frontend topbar and pricing page should read the reconciled account status."],
        ],
        [2200, 7160],
    )

    doc.add_heading("5. Billing Flow", level=1)
    add_table(
        doc,
        ["Job", "Endpoint", "Reason"],
        [
            ["First paid subscription", "POST /api/billing/checkout-session", "Stripe Checkout owns first purchase and payment collection."],
            ["Billing management", "POST /api/billing/portal-session", "Stripe Portal owns payment methods, invoices, cancellation, and downgrades."],
            ["Creator to Pro upgrade", "POST /api/billing/subscription-plan", "The app replaces one existing subscription item price in place."],
            ["Provider reconciliation", "POST /api/billing/webhook", "Stripe webhook is the source of truth for paid access."],
        ],
        [2350, 2850, 4160],
    )
    add_code_block(
        doc,
        [
            "PricingPage Pro button",
            "-> POST /api/billing/subscription-plan { plan_code: 'pro' }",
            "-> billing_service.change_subscription_plan",
            "-> load current provider subscription from PostgreSQL",
            "-> retrieve Stripe subscription and existing subscription item",
            "-> Stripe Subscription.modify(items=[{ id, price: STRIPE_PRO_PRICE_ID }])",
            "-> sync updated subscription into local database",
            "-> return account status",
        ],
    )
    add_callout(
        doc,
        "Billing safety decision",
        "The direct upgrade path changes the existing subscription item. It does not create a new Checkout session for an already-subscribed user, which avoids duplicate subscriptions and confusing invoices.",
    )

    doc.add_heading("6. Runtime Modes", level=1)
    add_table(
        doc,
        ["Mode", "Core environment", "What it optimizes"],
        [
            ["Production lightweight", "SONICMIND_MODE=production_light; SONICMIND_RETRIEVAL_BACKEND=lexical; ENABLE_LOCAL_EMBEDDING_MODEL=false; ENABLE_RERANKER=false", "Render 2 GB stability. Avoids torch, sentence-transformers, FAISS, and startup model loading."],
            ["Local semantic", "SONICMIND_MODE=local_semantic; SONICMIND_RETRIEVAL_BACKEND=faiss; ENABLE_LOCAL_EMBEDDING_MODEL=true; ENABLE_RERANKER=true", "Higher local answer quality and semantic retrieval testing with heavier dependencies."],
        ],
        [2100, 4300, 2960],
    )
    add_paragraph(
        doc,
        "Production mode intentionally sacrifices some semantic recall so the deployed backend can stay stable under the Render 2 GB memory ceiling. Semantic mode remains available for local testing or a future larger backend instance.",
    )

    doc.add_heading("7. Retrieval and Music Reasoning", level=1)
    add_paragraph(
        doc,
        "The music pipeline is evidence-first. It classifies the query, gathers local or trusted-source evidence, normalizes entities, then synthesizes the answer. Spotify cards are resolved from specific candidates rather than broad query text.",
    )
    add_bullet_list(
        doc,
        [
            "Local knowledge is attempted first when it is relevant.",
            "Trusted music sources and web retrieval help with artist profiles, current recommendations, and newly popular tracks.",
            "Spotify card generation should follow the resolved artist, album, or track candidates used in the answer.",
            "The production lexical mode is stable but can be less accurate than local semantic retrieval for niche artists and current-context questions.",
        ],
    )

    doc.add_heading("8. Frontend State and UI Boundaries", level=1)
    add_table(
        doc,
        ["File", "What to inspect first"],
        [
            ["frontend/src/App.jsx", "Top-level shell, routes, topbar plan/user display."],
            ["frontend/src/pages/PricingPage.jsx", "Plan cards, checkout, portal session, and direct Pro upgrade behavior."],
            ["frontend/src/pages/ChatPage.jsx", "Chat input, message list, account status, sources, favorites, and history panels."],
            ["frontend/src/api/client.js", "Backend URL selection, bearer token injection, and request helpers."],
            ["frontend/src/store/authStore.js", "Persisted token, user, and conversation state."],
        ],
        [3300, 6060],
    )
    add_paragraph(
        doc,
        "React Query owns server state. Zustand owns persisted client session state. After billing changes, the pricing page updates both React Query cache and stored user state so the topbar and chat quota display do not lag behind.",
    )

    doc.add_heading("9. Environment and Secret Boundaries", level=1)
    add_callout(
        doc,
        "Do not expose backend secrets",
        "Only VITE_API_BASE_URL belongs in the frontend bundle. Database URLs, Stripe keys, OpenAI keys, JWT secrets, Spotify secrets, Tavily keys, and Discogs tokens must stay backend-only.",
    )
    add_table(
        doc,
        ["Variable group", "Examples", "Where it belongs"],
        [
            ["Backend secrets", "DATABASE_URL, BACKEND_SECRET_KEY, STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, OPENAI_API_KEY", "Render backend environment only."],
            ["Backend public behavior", "APP_ENV, SONICMIND_MODE, SONICMIND_RETRIEVAL_BACKEND, RAG_TOP_K", "Render backend environment and local .env files."],
            ["Frontend config", "VITE_API_BASE_URL=https://sonicmind-api.onrender.com", "Render frontend static site environment only."],
        ],
        [2200, 4550, 2610],
    )

    doc.add_heading("10. Verification Commands", level=1)
    add_code_block(
        doc,
        [
            ".venv/bin/python -m pytest",
            "cd frontend && npm run build",
            ".venv/bin/python scripts/memory_probe.py --mode lexical",
            "cd frontend && FRONTEND_URL=https://sonicmind.onrender.com BACKEND_URL=https://sonicmind-api.onrender.com EXPECTED_RETRIEVAL_BACKEND=lexical npm run test:prod",
        ],
    )
    add_table(
        doc,
        ["Test group", "Why it matters"],
        [
            ["tests/test_api.py", "Health, chat API, and account status behavior."],
            ["tests/test_billing_api.py", "Checkout, portal, subscription-plan, and webhook routes."],
            ["tests/test_billing_service.py", "Stripe sync and plan-change service behavior."],
            ["tests/test_runtime_modes.py", "Lightweight import behavior and semantic mode configuration."],
            ["tests/test_music_intent.py", "Music query understanding and route selection."],
            ["tests/test_spotify_artist_albums.py", "Spotify artist, album, and fallback behavior."],
        ],
        [3350, 6010],
    )

    doc.add_heading("11. Review Notes and Remaining Risks", level=1)
    add_bullet_list(
        doc,
        [
            "README and rationale have been updated from older Streamlit wording to the current React + FastAPI deployment.",
            "render.yaml now points to https://sonicmind.onrender.com instead of the older frontend placeholder.",
            "Stripe billing docs now include direct Creator to Pro upgrades and the production webhook URL.",
            "Extra-pack purchases remain planned rather than implemented.",
            "A future larger launch should harden auth/session management beyond the current prototype token.",
        ],
    )

    doc.add_heading("12. Recommended Reading Order", level=1)
    add_numbered_list(
        doc,
        [
            "README.md for setup, deployed URLs, and developer commands.",
            "README_TECHNICAL_RATIONALE.md for architecture decisions and debugging history.",
            "backend/main.py and backend/schemas.py for API boundaries.",
            "backend/services/chat_service.py and src/rag_pipeline.py for chat behavior.",
            "src/services/quota_service.py for usage enforcement.",
            "backend/services/billing_service.py for Stripe behavior.",
            "frontend/src/pages/PricingPage.jsx and frontend/src/pages/ChatPage.jsx for user workflows.",
            "Run the test suite before changing production-facing behavior.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
